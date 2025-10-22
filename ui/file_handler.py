import streamlit as st
import base64
import mimetypes
import os
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Any
from PIL import Image
import io
import PyPDF2
import docx
import pandas as pd
import json
import yaml
from datetime import datetime

class FileHandler:
    """Handle all file types like Claude - images, PDFs, documents, code, etc."""
    
    # Supported file extensions matching Claude
    SUPPORTED_EXTENSIONS = {
        # Images
        'image': ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.svg', '.webp', '.ico'],
        # Documents
        'document': ['.pdf', '.doc', '.docx', '.txt', '.md', '.rtf', '.odt'],
        # Code files
        'code': ['.py', '.js', '.ts', '.jsx', '.tsx', '.java', '.c', '.cpp', '.h', 
                '.cs', '.php', '.rb', '.go', '.rs', '.kt', '.swift', '.m', '.scala',
                '.r', '.jl', '.lua', '.pl', '.sh', '.bash', '.zsh', '.fish', '.ps1',
                '.html', '.css', '.scss', '.sass', '.less', '.xml', '.json', '.yaml',
                '.yml', '.toml', '.ini', '.cfg', '.conf', '.sql', '.dockerfile', '.vue',
                '.dart', '.elm', '.clj', '.ex', '.erl', '.hs', '.ml', '.fs', '.nim'],
        # Data files
        'data': ['.csv', '.tsv', '.json', '.xml', '.yaml', '.yml', '.sqlite', '.db'],
        # Other
        'other': []
    }
    
    # Max file sizes (matching Claude's limits)
    MAX_FILE_SIZES = {
        'image': 10 * 1024 * 1024,      # 10MB for images
        'document': 50 * 1024 * 1024,   # 50MB for documents
        'code': 10 * 1024 * 1024,       # 10MB for code
        'data': 50 * 1024 * 1024,       # 50MB for data
        'default': 100 * 1024 * 1024    # 100MB default
    }
    
    @classmethod
    def get_file_type(cls, filename: str) -> str:
        """Determine file type from extension"""
        ext = Path(filename).suffix.lower()
        
        for file_type, extensions in cls.SUPPORTED_EXTENSIONS.items():
            if ext in extensions:
                return file_type
        return 'other'
    
    @classmethod
    def is_supported(cls, filename: str) -> bool:
        """Check if file type is supported"""
        ext = Path(filename).suffix.lower()
        for extensions in cls.SUPPORTED_EXTENSIONS.values():
            if ext in extensions:
                return True
        return False
    
    @classmethod
    def validate_file(cls, file) -> Tuple[bool, str]:
        """Validate uploaded file"""
        if not file:
            return False, "No file provided"
            
        filename = file.name
        if not cls.is_supported(filename):
            return False, f"File type not supported: {Path(filename).suffix}"
            
        file_type = cls.get_file_type(filename)
        max_size = cls.MAX_FILE_SIZES.get(file_type, cls.MAX_FILE_SIZES['default'])
        
        if file.size > max_size:
            return False, f"File too large: {file.size / (1024*1024):.2f}MB (max: {max_size / (1024*1024):.2f}MB)"
            
        return True, "Valid"
    
    @classmethod
    def process_file(cls, file) -> Dict[str, Any]:
        """Process uploaded file and extract content"""
        filename = file.name
        file_type = cls.get_file_type(filename)
        file_size = file.size
        
        result = {
            'name': filename,
            'type': file_type,
            'size': file_size,
            'extension': Path(filename).suffix.lower(),
            'content': None,
            'preview': None,
            'metadata': {},
            'error': None
        }
        
        try:
            if file_type == 'image':
                result.update(cls._process_image(file))
            elif file_type == 'document':
                result.update(cls._process_document(file))
            elif file_type == 'code':
                result.update(cls._process_code(file))
            elif file_type == 'data':
                result.update(cls._process_data(file))
            else:
                result['content'] = file.getvalue()
                
        except Exception as e:
            result['error'] = str(e)
            
        return result
    
    @classmethod
    def _process_image(cls, file) -> Dict:
        """Process image files"""
        try:
            image = Image.open(file)
            
            # Generate preview
            preview_size = (300, 300)
            image.thumbnail(preview_size, Image.Resampling.LANCZOS)
            
            # Convert to base64 for display
            buffer = io.BytesIO()
            image.save(buffer, format=image.format or 'PNG')
            img_str = base64.b64encode(buffer.getvalue()).decode()
            
            return {
                'content': file.getvalue(),
                'preview': f"data:image/png;base64,{img_str}",
                'metadata': {
                    'width': image.width,
                    'height': image.height,
                    'mode': image.mode,
                    'format': image.format
                }
            }
        except Exception as e:
            return {'error': f"Failed to process image: {e}"}
    
    @classmethod
    def _process_document(cls, file) -> Dict:
        """Process document files"""
        filename = file.name
        ext = Path(filename).suffix.lower()
        
        try:
            if ext == '.pdf':
                return cls._process_pdf(file)
            elif ext in ['.doc', '.docx']:
                return cls._process_word(file)
            elif ext in ['.txt', '.md', '.rtf']:
                content = file.getvalue().decode('utf-8', errors='ignore')
                return {
                    'content': content,
                    'preview': content[:1000] + '...' if len(content) > 1000 else content
                }
            else:
                return {'content': file.getvalue()}
                
        except Exception as e:
            return {'error': f"Failed to process document: {e}"}
    
    @classmethod
    def _process_pdf(cls, file) -> Dict:
        """Process PDF files"""
        try:
            pdf_reader = PyPDF2.PdfReader(file)
            num_pages = len(pdf_reader.pages)
            
            # Extract text from all pages
            text_content = []
            for page_num in range(min(num_pages, 50)):  # Limit to first 50 pages
                page = pdf_reader.pages[page_num]
                text_content.append(page.extract_text())
                
            full_text = '\n\n'.join(text_content)
            
            return {
                'content': file.getvalue(),
                'preview': full_text[:2000] + '...' if len(full_text) > 2000 else full_text,
                'metadata': {
                    'pages': num_pages,
                    'text_content': full_text
                }
            }
        except Exception as e:
            return {'error': f"Failed to process PDF: {e}"}
    
    @classmethod
    def _process_word(cls, file) -> Dict:
        """Process Word documents"""
        try:
            doc = docx.Document(file)
            full_text = []
            
            for paragraph in doc.paragraphs:
                full_text.append(paragraph.text)
                
            text_content = '\n'.join(full_text)
            
            return {
                'content': file.getvalue(),
                'preview': text_content[:2000] + '...' if len(text_content) > 2000 else text_content,
                'metadata': {
                    'paragraphs': len(doc.paragraphs),
                    'text_content': text_content
                }
            }
        except Exception as e:
            return {'error': f"Failed to process Word document: {e}"}
    
    @classmethod
    def _process_code(cls, file) -> Dict:
        """Process code files"""
        try:
            content = file.getvalue().decode('utf-8', errors='ignore')
            
            # Count lines, detect language
            lines = content.split('\n')
            
            return {
                'content': content,
                'preview': content[:2000] + '...' if len(content) > 2000 else content,
                'metadata': {
                    'lines': len(lines),
                    'language': cls._detect_language(file.name)
                }
            }
        except Exception as e:
            return {'error': f"Failed to process code file: {e}"}
    
    @classmethod
    def _process_data(cls, file) -> Dict:
        """Process data files"""
        filename = file.name
        ext = Path(filename).suffix.lower()
        
        try:
            if ext in ['.csv', '.tsv']:
                delimiter = '\t' if ext == '.tsv' else ','
                df = pd.read_csv(file, delimiter=delimiter, nrows=100)
                
                return {
                    'content': file.getvalue(),
                    'preview': df.to_string(max_rows=10),
                    'metadata': {
                        'rows': len(df),
                        'columns': list(df.columns),
                        'shape': df.shape,
                        'df_head': df.head().to_dict()
                    }
                }
            elif ext in ['.json']:
                content = file.getvalue().decode('utf-8')
                data = json.loads(content)
                
                return {
                    'content': content,
                    'preview': json.dumps(data, indent=2)[:2000],
                    'metadata': {
                        'keys': list(data.keys()) if isinstance(data, dict) else None,
                        'type': type(data).__name__
                    }
                }
            elif ext in ['.yaml', '.yml']:
                content = file.getvalue().decode('utf-8')
                data = yaml.safe_load(content)
                
                return {
                    'content': content,
                    'preview': yaml.dump(data)[:2000],
                    'metadata': {
                        'keys': list(data.keys()) if isinstance(data, dict) else None
                    }
                }
            else:
                return {'content': file.getvalue()}
                
        except Exception as e:
            return {'error': f"Failed to process data file: {e}"}
    
    @classmethod
    def _detect_language(cls, filename: str) -> str:
        """Detect programming language from file extension"""
        ext_to_lang = {
            '.py': 'python', '.js': 'javascript', '.ts': 'typescript',
            '.java': 'java', '.c': 'c', '.cpp': 'cpp', '.cs': 'csharp',
            '.rb': 'ruby', '.go': 'go', '.rs': 'rust', '.php': 'php',
            '.swift': 'swift', '.kt': 'kotlin', '.r': 'r', '.jl': 'julia',
            '.sh': 'bash', '.sql': 'sql', '.html': 'html', '.css': 'css',
            '.json': 'json', '.yaml': 'yaml', '.yml': 'yaml', '.xml': 'xml'
        }
        
        ext = Path(filename).suffix.lower()
        return ext_to_lang.get(ext, 'text')

def render_file_upload_zone(key_suffix=""):
    """Render drag-and-drop file upload zone like Claude with better multi-file support"""
    
    st.markdown("""
        <style>
        .uploadzone {
            border: 2px dashed #4CAF50;
            border-radius: 10px;
            padding: 30px;
            text-align: center;
            background: linear-gradient(135deg, #f5f7fa 0%, #c3cfe2 100%);
            margin-bottom: 20px;
            transition: all 0.3s ease;
        }
        .uploadzone:hover {
            border-color: #45a049;
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            transform: scale(1.02);
        }
        .file-info {
            background: white;
            padding: 10px;
            border-radius: 5px;
            margin: 10px 0;
            box-shadow: 0 2px 4px rgba(0,0,0,0.1);
        }
        .file-counter {
            background: #4CAF50;
            color: white;
            padding: 8px 16px;
            border-radius: 20px;
            display: inline-block;
            margin: 10px 0;
            font-weight: bold;
        }
        </style>
    """, unsafe_allow_html=True)
    
    st.markdown('<div class="uploadzone">', unsafe_allow_html=True)
    
    # Clear instructions for multiple file upload
    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        st.markdown("### 📂 Upload Multiple Files")
        st.markdown("**Drag & Drop** or **Click to Browse**")
        st.markdown("💡 **Tip:** Use `Ctrl+A` to select all files in a folder!")
    
    st.markdown('</div>', unsafe_allow_html=True)
    
    # File uploader with clear multiple files support
    uploaded_files = st.file_uploader(
        "Choose files",
        accept_multiple_files=True,
        help="✅ Upload unlimited files at once • ✅ All file types supported • ✅ Select multiple with Ctrl/Cmd+Click • ✅ Direct folder upload",
        key=f"multi_file_upload_{key_suffix}",
        type=None  # Accept all file types
    )
    
    # Show file counter and list
    if uploaded_files:
        st.markdown(f'<div class="file-counter">✅ {len(uploaded_files)} file(s) ready to process</div>', unsafe_allow_html=True)
        
        # Show file list with sizes
        with st.expander(f"📋 Selected Files ({len(uploaded_files)})", expanded=True):
            for i, file in enumerate(uploaded_files, 1):
                file_size_mb = file.size / (1024 * 1024)
                icon = "📄" if file.name.endswith(('.txt', '.md', '.py', '.js')) else "📎"
                st.markdown(f"{i}. {icon} **{file.name}** ({file_size_mb:.2f} MB)")
    
    processed_files = []
    
    if uploaded_files:
        # Add progress bar for multiple files
        if len(uploaded_files) > 1:
            progress_bar = st.progress(0)
            status_text = st.empty()
        
        for idx, file in enumerate(uploaded_files):
            valid, message = FileHandler.validate_file(file)
            
            if valid:
                # Update progress for multiple files
                if len(uploaded_files) > 1:
                    progress = (idx + 1) / len(uploaded_files)
                    progress_bar.progress(progress)
                    status_text.text(f"Processing {idx + 1}/{len(uploaded_files)}: {file.name}")
                
                with st.spinner(f"Processing {file.name}..."):
                    processed = FileHandler.process_file(file)
                    processed_files.append(processed)
                    
                    # Don't show individual previews if many files
                    if len(uploaded_files) <= 5:
                        # Show preview for small number of files
                        with st.expander(f"📄 {file.name} ({file.size / 1024:.1f} KB)", expanded=False):
                            if processed.get('error'):
                                st.error(processed['error'])
                            else:
                                file_type = processed['type']
                                
                                if file_type == 'image' and processed.get('preview'):
                                    st.image(processed['preview'])
                                elif processed.get('preview'):
                                    if file_type == 'code':
                                        st.code(processed['preview'], language=processed['metadata'].get('language', 'text'))
                                    else:
                                        st.text(processed['preview'])
                                        
                                if processed.get('metadata'):
                                    st.json(processed['metadata'])
            else:
                st.error(f"❌ {file.name}: {message}")
        
        # Clear progress bar after completion
        if len(uploaded_files) > 1:
            progress_bar.empty()
            status_text.empty()
            st.success(f"✅ Successfully processed {len(processed_files)} files!")
                
    return processed_files

def create_file_reference(file_info: Dict) -> str:
    """Create an inline reference to a file in chat"""
    return f"@{file_info['name']}"

def parse_file_references(text: str, project_files: Dict) -> List[str]:
    """Parse @filename references in text"""
    import re
    
    # Find all @filename references
    pattern = r'@([\w\-\.]+(?:\.\w+)?)'
    matches = re.findall(pattern, text)
    
    referenced_files = []
    for match in matches:
        # Check if file exists in project
        for file_path in project_files.keys():
            if match in file_path or Path(file_path).name == match:
                referenced_files.append(file_path)
                break
                
    return referenced_files

def render_file_in_chat(file_info: Dict):
    """Render a file reference in chat like Claude"""
    file_type = file_info.get('type', 'other')
    
    # Create a compact file card
    col1, col2 = st.columns([1, 4])
    
    with col1:
        # File type icon
        icons = {
            'image': '🖼️',
            'document': '📄',
            'code': '💻',
            'data': '📊',
            'other': '📎'
        }
        st.markdown(f"### {icons.get(file_type, '📎')}")
        
    with col2:
        st.markdown(f"**{file_info['name']}**")
        if file_info.get('metadata'):
            if file_type == 'image':
                st.caption(f"{file_info['metadata'].get('width')}x{file_info['metadata'].get('height')} pixels")
            elif file_type == 'code':
                st.caption(f"{file_info['metadata'].get('lines')} lines • {file_info['metadata'].get('language')}")
            elif file_type == 'document':
                if 'pages' in file_info['metadata']:
                    st.caption(f"{file_info['metadata']['pages']} pages")
            elif file_type == 'data':
                if 'shape' in file_info['metadata']:
                    st.caption(f"{file_info['metadata']['shape'][0]} rows × {file_info['metadata']['shape'][1]} columns")

class FileSearch:
    """Search and filter files in project"""
    
    @staticmethod
    def search_files(files: Dict, query: str, file_types: List[str] = None) -> Dict:
        """Search files by name and content"""
        if not query and not file_types:
            return files
            
        filtered = {}
        query_lower = query.lower() if query else ""
        
        for file_path, file_obj in files.items():
            # Filter by type
            if file_types:
                file_type = FileHandler.get_file_type(file_path)
                if file_type not in file_types:
                    continue
                    
            # Filter by query
            if query:
                # Check filename
                if query_lower in file_path.lower():
                    filtered[file_path] = file_obj
                    continue
                    
                # Check content for text files
                if hasattr(file_obj, 'content') and file_obj.content:
                    if query_lower in str(file_obj.content).lower():
                        filtered[file_path] = file_obj
                        continue
            else:
                filtered[file_path] = file_obj
                
        return filtered
    
    @staticmethod
    def render_search_bar():
        """Render file search interface"""
        col1, col2 = st.columns([2, 1])
        
        with col1:
            search_query = st.text_input("🔍 Search files", placeholder="Search by name or content...")
            
        with col2:
            file_types = st.multiselect(
                "Filter by type",
                options=['image', 'document', 'code', 'data', 'other'],
                default=None
            )
            
        return search_query, file_types